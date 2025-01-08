import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext
import pandas as pd
import numpy as np
from sklearn.ensemble import RandomForestRegressor
import threading
from docx import Document
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg


class RetentionPredictionApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Audience Retention Prediction")
        self.root.geometry("1200x800")

        self.model = None
        self.data = {}

        self.create_gui()

    def create_gui(self):
        self.control_frame = ttk.LabelFrame(self.root, text="Controls", padding="10")
        self.control_frame.pack(fill="x", padx=10, pady=5)

        self.output_frame = ttk.Notebook(self.root)
        self.output_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.create_output_elements()
        self.create_control_elements()

        self.log("Instructions:")
        self.log(
            "1. Use 'Load Data Folder' to upload the folder containing video audience retention data (.xlsx) and transcripts (.docx).")
        self.log("2. Use 'Load Sample File' to upload the sample .xlsx file with missing values for prediction.")
        self.log(
            "3. Use 'Predict Retention' to generate predictions for the missing retention values (positions 0 to 5) and view the results in the Log and Plot tabs.")

    def create_control_elements(self):
        ttk.Button(self.control_frame, text="Load Data Folder",
                   command=self.load_data_folder).pack(side="left", padx=5)

        ttk.Button(self.control_frame, text="Load Sample File",
                   command=self.load_sample_file).pack(side="left", padx=5)

        self.predict_button = ttk.Button(self.control_frame, text="Predict Retention",
                                         command=self.predict_retention, state="disabled")
        self.predict_button.pack(side="left", padx=5)

    def create_output_elements(self):
        # Log tab
        self.log_frame = ttk.Frame(self.output_frame)
        self.output_frame.add(self.log_frame, text="Log")
        self.log_text = scrolledtext.ScrolledText(self.log_frame, height=20)
        self.log_text.pack(fill="both", expand=True)

        # Plot tab
        self.plot_frame = ttk.Frame(self.output_frame)
        self.output_frame.add(self.plot_frame, text="Plot")

    def log(self, message):
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)

    def load_data_folder(self):
        folder_path = filedialog.askdirectory()
        if not folder_path:
            self.log("No folder selected.")
            return

        self.log(f"Loading data from folder: {folder_path}")

        # Simulate loading data from folder
        # Assume each pair of .xlsx and .docx is named consistently, e.g., "video_a.xlsx" and "video_a.docx"
        self.data = {}  # Reset data
        import os

        for file in os.listdir(folder_path):
            if file.endswith(".xlsx"):
                base_name = file.replace(".xlsx", "")
                xlsx_path = os.path.join(folder_path, file)
                docx_path = os.path.join(folder_path, base_name + ".docx")

                if os.path.exists(docx_path):
                    self.data[base_name] = {
                        "xlsx": xlsx_path,
                        "docx": docx_path
                    }

        self.log(f"Loaded {len(self.data)} video data pairs.")

    def load_sample_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if not file_path:
            self.log("No sample file selected.")
            return

        self.sample_file = file_path
        self.log(f"Sample file loaded: {file_path}")
        self.predict_button.config(state="normal")

    def extract_features(self, xlsx_path, docx_path):
        # Load retention data from Excel
        retention_data = pd.read_excel(xlsx_path)
        retention_values = retention_data['Absolute audience retention (%)'].apply(
            lambda x: x if isinstance(x, (int, float)) else np.nan
        ).dropna().values

        # Load transcript from DOCX
        doc = Document(docx_path)
        transcript = " ".join([para.text for para in doc.paragraphs])

        # Example: Compute basic features (you can extend this)
        avg_retention = np.mean(retention_values)
        word_count = len(transcript.split())

        return [avg_retention, word_count]

    def train_model(self):
        self.log("Training model...")

        X = []
        y = []

        for base_name, paths in self.data.items():
            features = self.extract_features(paths['xlsx'], paths['docx'])
            retention_data = pd.read_excel(paths['xlsx'])
            retention_values = retention_data['Absolute audience retention (%)'].apply(
                lambda x: x if isinstance(x, (int, float)) else np.nan
            ).dropna().values

            for i, retention in enumerate(retention_values):
                X.append(features + [i])  # Add position as a feature
                y.append(retention)

        X = np.array(X)
        y = np.array(y)

        self.model = RandomForestRegressor(n_estimators=100, random_state=42)
        self.model.fit(X, y)

        self.log("Model training completed.")

    def predict_retention(self):
        if not self.model:
            self.train_model()

        sample_data = pd.read_excel(self.sample_file)
        sample_transcript_path = filedialog.askopenfilename(filetypes=[("DOCX files", "*.docx")])
        if not sample_transcript_path:
            self.log("No transcript for sample file selected.")
            return

        doc = Document(sample_transcript_path)
        sample_transcript = " ".join([para.text for para in doc.paragraphs])

        avg_retention = sample_data['Absolute audience retention (%)'].apply(
            lambda x: x if isinstance(x, (int, float)) else np.nan
        ).dropna().mean()
        word_count = len(sample_transcript.split())

        predictions = []
        for i in range(min(6, len(sample_data))):  # Predict for positions 0 to 5 or all available positions
            features = [avg_retention, word_count, i]
            predicted_retention = self.model.predict([features])[0]
            predictions.append(predicted_retention)

        self.log("Predicted Retention Values:")
        for i, value in enumerate(predictions):
            self.log(f"Position {i}: {value:.2f}%")

        # Plot predictions
        self.plot_predictions(predictions)

    def plot_predictions(self, predictions):
        for widget in self.plot_frame.winfo_children():
            widget.destroy()

        fig, ax = plt.subplots(figsize=(8, 5))
        positions = list(range(len(predictions)))
        ax.plot(positions, predictions, marker='o', linestyle='-', color='b')
        ax.set_title("Predicted Audience Retention")
        ax.set_xlabel("Position (0-5)")
        ax.set_ylabel("Retention (%)")
        ax.grid(True)

        canvas = FigureCanvasTkAgg(fig, master=self.plot_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)


if __name__ == "__main__":
    root = tk.Tk()
    app = RetentionPredictionApp(root)
    root.mainloop()
